import asyncio
import hashlib
import logging
import time
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import aiofiles
from concurrent.futures import ThreadPoolExecutor
import magic
import json

import pdfplumber
from docx import Document as DocxDocument
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from collections import Counter
import re

from ..core.config import settings
from ..models.schemas import DocumentChunk, ProcessedDocument, DocumentMetadata, DocumentChunkMetadata
from ..core.exceptions import DocumentProcessingError, ValidationError
from ..utils.text_utils import TextProcessor, SemanticChunker

logger = logging.getLogger(__name__)


class AdvancedDocumentProcessor:
    """Advanced document processor with semantic chunking and metadata extraction."""

    def __init__(self):
        self.text_processor = TextProcessor()
        self.semantic_chunker = SemanticChunker()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.encoding = tiktoken.get_encoding("cl100k_base")

    async def process_document(
            self,
            file_path: Path,
            filename: str,
            use_semantic_chunking: bool = True,
            extract_metadata: bool = True
    ) -> ProcessedDocument:
        """Process a document with advanced features."""
        start_time = time.time()

        try:
            # Validate file
            await self._validate_file(file_path, filename)

            # Extract content and metadata
            content, base_metadata = await self._extract_content_and_metadata(file_path, filename)

            if not content.strip():
                raise DocumentProcessingError("Document contains no extractable text")

            # Generate content hash
            content_hash = self._generate_content_hash(content)

            # Clean and preprocess content
            cleaned_content = await self._clean_and_preprocess_content(content)

            # Perform chunking
            if use_semantic_chunking and len(cleaned_content) > 2000:
                chunks = await self._semantic_chunk_text(cleaned_content)
            else:
                chunks = await self._standard_chunk_text(cleaned_content)

            # Limit chunks to prevent memory issues
            if len(chunks) > settings.MAX_CHUNKS_PER_DOCUMENT:
                logger.warning(f"Document has {len(chunks)} chunks, limiting to {settings.MAX_CHUNKS_PER_DOCUMENT}")
                chunks = chunks[:settings.MAX_CHUNKS_PER_DOCUMENT]

            # Enhance chunks with metadata
            enhanced_chunks = await self._enhance_chunks_with_metadata(
                chunks, base_metadata, filename, content_hash
            )

            # Extract additional metadata
            if extract_metadata:
                additional_metadata = await self._extract_advanced_metadata(cleaned_content)
                base_metadata.update(additional_metadata)

            # Generate summary
            summary = await self._generate_document_summary(cleaned_content)

            # Extract key topics
            key_topics = await self._extract_key_topics(cleaned_content)

            processing_time = time.time() - start_time
            base_metadata["processing_time"] = processing_time

            return ProcessedDocument(
                filename=filename,
                content_hash=content_hash,
                total_chunks=len(enhanced_chunks),
                chunks=enhanced_chunks,
                metadata=DocumentMetadata(**base_metadata),
                summary=summary,
                word_count=len(cleaned_content.split()),
                processing_time=processing_time,
                key_topics=key_topics
            )

        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    async def _validate_file(self, file_path: Path, filename: str):
        """Validate file before processing."""
        if not file_path.exists():
            raise ValidationError(f"File not found: {file_path}")

        # Check file size
        file_size = file_path.stat().st_size
        if file_size > settings.MAX_FILE_SIZE:
            raise ValidationError(f"File size {file_size} exceeds maximum {settings.MAX_FILE_SIZE}")

        # Check file type
        file_extension = file_path.suffix.lower().lstrip('.')
        if file_extension not in settings.SUPPORTED_FORMATS:
            raise ValidationError(f"Unsupported file format: {file_extension}")

        # Verify MIME type
        mime_type = magic.from_file(str(file_path), mime=True)
        expected_mime_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'md': 'text/plain',
            'html': 'text/html'
        }

        expected_mime = expected_mime_types.get(file_extension)
        if expected_mime and not mime_type.startswith(expected_mime.split('/')[0]):
            logger.warning(f"MIME type mismatch for {filename}: expected {expected_mime}, got {mime_type}")

    async def _extract_content_and_metadata(
            self,
            file_path: Path,
            filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract content and metadata from various file formats."""
        file_extension = file_path.suffix.lower()

        if file_extension == '.pdf':
            return await self._extract_pdf_content(file_path)
        elif file_extension == '.docx':
            return await self._extract_docx_content(file_path)
        elif file_extension in ['.txt', '.md']:
            return await self._extract_text_content(file_path)
        elif file_extension == '.html':
            return await self._extract_html_content(file_path)
        else:
            raise DocumentProcessingError(f"Unsupported file format: {file_extension}")

    async def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Advanced PDF extraction with comprehensive metadata."""

        def extract_pdf():
            content = ""
            metadata = {
                "pages": 0,
                "tables": [],
                "images": 0,
                "fonts": set(),
                "page_layouts": [],
                "text_blocks": 0,
                "has_forms": False,
                "is_scanned": False
            }

            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)

                    # Extract PDF metadata
                    if pdf.metadata:
                        metadata.update({
                            "title": pdf.metadata.get("Title", ""),
                            "author": pdf.metadata.get("Author", ""),
                            "subject": pdf.metadata.get("Subject", ""),
                            "creator": pdf.metadata.get("Creator", ""),
                            "producer": pdf.metadata.get("Producer", ""),
                            "creation_date": str(pdf.metadata.get("CreationDate", "")),
                            "modification_date": str(pdf.metadata.get("ModDate", ""))
                        })

                    total_chars = 0
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        page_text = page.extract_text() or ""
                        total_chars += len(page_text)

                        if page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                            metadata["text_blocks"] += len(page_text.split('\n\n'))

                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables):
                                if table and len(table) > 0:
                                    metadata["tables"].append({
                                        "page": page_num + 1,
                                        "table_id": f"table_{page_num}_{table_idx}",
                                        "rows": len(table),
                                        "cols": len(table[0]) if table[0] else 0
                                    })

                                    # Convert table to text
                                    table_text = self._table_to_text(table)
                                    content += f"\n[TABLE {len(metadata['tables'])}]\n{table_text}\n"

                        # Extract images
                        if hasattr(page, 'images') and page.images:
                            metadata["images"] += len(page.images)

                        # Extract fonts
                        if hasattr(page, 'chars') and page.chars:
                            page_fonts = {char.get('fontname', 'Unknown') for char in page.chars}
                            metadata["fonts"].update(page_fonts)

                        # Check for forms
                        if hasattr(page, 'annots') and page.annots:
                            metadata["has_forms"] = True

                        # Analyze page layout
                        layout_info = {
                            "page": page_num + 1,
                            "width": float(page.width) if page.width else 0,
                            "height": float(page.height) if page.height else 0,
                            "text_density": len(page_text) / (
                                        page.width * page.height) if page.width and page.height else 0
                        }
                        metadata["page_layouts"].append(layout_info)

                    # Determine if document is likely scanned
                    avg_chars_per_page = total_chars / len(pdf.pages) if pdf.pages else 0
                    metadata["is_scanned"] = avg_chars_per_page < 100  # Heuristic

                    metadata["fonts"] = list(metadata["fonts"])

            except Exception as e:
                logger.error(f"Error extracting PDF content: {str(e)}")
                raise DocumentProcessingError(f"PDF extraction failed: {str(e)}")

            return content, metadata

        return await asyncio.get_event_loop().run_in_executor(
            self.executor, extract_pdf
        )

    async def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX files with structure preservation."""

        def extract_docx():
            try:
                doc = DocxDocument(file_path)
                content = ""
                metadata = {
                    "paragraphs": 0,
                    "tables": 0,
                    "images": 0,
                    "styles": set(),
                    "headers": 0,
                    "footers": 0,
                    "sections": len(doc.sections)
                }

                # Extract core properties
                if doc.core_properties:
                    metadata.update({
                        "title": doc.core_properties.title or "",
                        "author": doc.core_properties.author or "",
                        "subject": doc.core_properties.subject or "",
                        "keywords": doc.core_properties.keywords or "",
                        "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                        "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                    })

                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                        metadata["paragraphs"] += 1

                        if paragraph.style:
                            metadata["styles"].add(paragraph.style.name)

                            # Count headers
                            if "heading" in paragraph.style.name.lower():
                                metadata["headers"] += 1

                # Extract tables
                for table_idx, table in enumerate(doc.tables):
                    metadata["tables"] += 1
                    table_text = ""

                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            table_text += row_text + "\n"

                    if table_text.strip():
                        content += f"\n[TABLE {metadata['tables']}]\n{table_text}\n"

                # Count images and other elements
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        metadata["images"] += 1

                # Extract headers and footers
                for section in doc.sections:
                    if section.header:
                        header_text = ""
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                header_text += paragraph.text + " "
                        if header_text.strip():
                            content = f"[HEADER] {header_text.strip()}\n" + content
                            metadata["headers"] += 1

                    if section.footer:
                        footer_text = ""
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                footer_text += paragraph.text + " "
                        if footer_text.strip():
                            content += f"\n[FOOTER] {footer_text.strip()}\n"
                            metadata["footers"] += 1

                metadata["styles"] = list(metadata["styles"])

            except Exception as e:
                logger.error(f"Error extracting DOCX content: {str(e)}")
                raise DocumentProcessingError(f"DOCX extraction failed: {str(e)}")

            return content, metadata

        return await asyncio.get_event_loop().run_in_executor(
            self.executor, extract_docx
        )

    async def _extract_text_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from text files with encoding detection."""
        try:
            # Try to detect encoding
            import chardet

            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'

            # Read with detected encoding
            async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                content = await file.read()

            metadata = {
                "lines": len(content.split('\n')),
                "characters": len(content),
                "encoding": encoding,
                "confidence": encoding_result.get('confidence', 0.0)
            }

            return content, metadata

        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")

    async def _extract_html_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from HTML files."""
        try:
            from bs4 import BeautifulSoup

            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)

            metadata = {
                "title": soup.title.string if soup.title else "",
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "tables": len(soup.find_all('table')),
                "headings": len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                "paragraphs": len(soup.find_all('p'))
            }

            return content, metadata

        except Exception as e:
            logger.error(f"Error extracting HTML content: {str(e)}")
            raise DocumentProcessingError(f"HTML extraction failed: {str(e)}")

    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table data to readable text format."""
        if not table:
            return ""

        # Filter out empty rows
        filtered_table = [row for row in table if any(cell and cell.strip() for cell in row)]

        if not filtered_table:
            return ""

        # Create formatted table
        text_lines = []
        for row in filtered_table:
            cleaned_row = [cell.strip() if cell else "" for cell in row]
            text_lines.append(" | ".join(cleaned_row))

        return "\n".join(text_lines)

    async def _clean_and_preprocess_content(self, content: str) -> str:
        """Clean and preprocess content."""
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r' +', ' ', content)

        # Remove page headers/footers patterns
        content = re.sub(r'\n--- Page \d+ ---\n', '\n', content)

        # Clean up table markers
        content = re.sub(r'\[TABLE \d+\]', '[TABLE]', content)

        # Remove control characters
        content = ''.join(char for char in content if ord(char) >= 32 or char in '\n\t')

        return content.strip()

    async def _semantic_chunk_text(self, content: str) -> List[str]:
        """Perform semantic chunking using embeddings."""
        try:
            chunks = await self.semantic_chunker.chunk_text(content)
            return chunks
        except Exception as e:
            logger.warning(f"Semantic chunking failed, falling back to standard chunking: {str(e)}")
            return await self._standard_chunk_text(content)

    async def _standard_chunk_text(self, content: str) -> List[str]:
        """Perform standard recursive chunking."""

        def chunk_text():
            docs = self.text_splitter.create_documents([content])
            return [doc.page_content for doc in docs]

        return await asyncio.get_event_loop().run_in_executor(
            self.executor, chunk_text
        )

    async def _enhance_chunks_with_metadata(
            self,
            chunks: List[str],
            metadata: Dict[str, Any],
            filename: str,
            content_hash: str
    ) -> List[DocumentChunk]:
        """Enhance chunks with rich metadata."""
        enhanced_chunks = []

        for idx, chunk in enumerate(chunks):
            # Generate chunk metadata
            chunk_metadata = DocumentChunkMetadata(
                chunk_id=f"{content_hash}_{idx}",
                filename=filename,
                chunk_index=idx,
                total_chunks=len(chunks),
                word_count=len(chunk.split()),
                char_count=len(chunk),
                source_metadata=metadata
            )

            # Estimate page number
            if "pages" in metadata and metadata["pages"] > 0:
                estimated_page = min(
                    int((idx / len(chunks)) * metadata["pages"]) + 1,
                    metadata["pages"]
                )
                chunk_metadata.estimated_page = estimated_page

            # Extract section title if available
            section_title = self._extract_section_title(chunk)
            if section_title:
                chunk_metadata.section_title = section_title

            enhanced_chunks.append(
                DocumentChunk(
                    content=chunk,
                    metadata=chunk_metadata,
                    embedding=None  # Will be populated by embedding service
                )
            )

        return enhanced_chunks

    def _extract_section_title(self, chunk: str) -> Optional[str]:
        """Extract section title from chunk if available."""
        lines = chunk.split('\n')
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            if line and (
                    line.isupper() or
                    re.match(r'^[A-Z][^.!?]*$', line) or
                    re.match(r'^\d+\.?\s+[A-Z]', line)
            ):
                return line[:100]  # Limit length
        return None

    async def _extract_advanced_metadata(self, content: str) -> Dict[str, Any]:
        """Extract advanced metadata using simple text analysis."""
        metadata = {}

        # Text statistics
        words = content.split()
        metadata.update({
            "unique_words": len(set(word.lower() for word in words)),
            "avg_word_length": sum(len(word) for word in words) / len(words) if words else 0,
            "sentences_estimated": len(re.findall(r'[.!?]+', content)),
            "paragraphs_estimated": len(content.split('\n\n'))
        })

        # Simple entity extraction using regex patterns
        entities = {
            "EMAIL": re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', content),
            "URL": re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', content),
            "MONEY": re.findall(r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:dollars?|USD|euros?|EUR)', content, re.IGNORECASE),
            "DATE": re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b', content),
            "PERSON": re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', content)
        }

        # Remove duplicates and limit results
        for key in entities:
            entities[key] = list(set(entities[key]))[:10]

        metadata["entities"] = entities

        return metadata

    async def _generate_document_summary(self, content: str) -> str:
        """Generate an intelligent document summary."""
        # Use extractive summarization
        sentences = self.text_processor.split_into_sentences(content)

        if len(sentences) <= 3:
            return content[:500] + "..." if len(content) > 500 else content

        # Score sentences based on position, length, and keyword frequency
        sentence_scores = []
        word_freq = Counter(word.lower() for word in content.split())

        for i, sentence in enumerate(sentences):
            score = 0
            words = sentence.split()

            # Position score (beginning and end are important)
            if i < 3:
                score += 2
            elif i >= len(sentences) - 3:
                score += 1

            # Length score (prefer medium-length sentences)
            if 10 <= len(words) <= 30:
                score += 1

            # Keyword frequency score
            for word in words:
                score += word_freq.get(word.lower(), 0) / len(words)

            sentence_scores.append((score, sentence))

        # Select top sentences
        sentence_scores.sort(reverse=True)
        selected_sentences = [sent for _, sent in sentence_scores[:5]]

        # Reorder by original position
        summary_sentences = []
        for sentence in sentences:
            if sentence in selected_sentences:
                summary_sentences.append(sentence)

        summary = " ".join(summary_sentences)

        # Ensure summary is within word limit
        words = summary.split()
        if len(words) > 150:
            summary = " ".join(words[:150]) + "..."

        return summary

    async def _extract_key_topics(self, content: str) -> List[str]:
        """Extract key topics using TF-IDF and NLP."""
        try:
            # Simple keyword extraction
            words = re.findall(r'\b[a-zA-Z]{4,}\b', content.lower())

            # Remove common stop words
            stop_words = {
                'this', 'that', 'with', 'have', 'will', 'from', 'they', 'know',
                'want', 'been', 'good', 'much', 'some', 'time', 'very', 'when',
                'come', 'here', 'just', 'like', 'long', 'make', 'many', 'over',
                'such', 'take', 'than', 'them', 'well', 'were', 'what', 'your',
                'about', 'after', 'again', 'before', 'being', 'could', 'every',
                'first', 'found', 'great', 'group', 'large', 'last', 'little',
                'most', 'never', 'number', 'only', 'other', 'people', 'place',
                'right', 'same', 'should', 'small', 'still', 'their', 'there',
                'these', 'those', 'through', 'under', 'until', 'water', 'where',
                'which', 'while', 'world', 'would', 'write', 'years', 'young'
            }

            filtered_words = [word for word in words if word not in stop_words and len(word) > 4]

            # Count frequency
            word_freq = Counter(filtered_words)

            # Get top keywords
            top_keywords = [word for word, _ in word_freq.most_common(15)]

            # Simple noun phrase extraction using regex
            noun_phrases = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', content[:5000])
            noun_phrases = [phrase.lower() for phrase in noun_phrases 
                           if len(phrase.split()) <= 3 and len(phrase) > 3]

                # Combine with keywords
                all_topics = list(set(top_keywords + noun_phrases))
                return all_topics[:10]

        except Exception as e:
            logger.warning(f"Topic extraction failed: {str(e)}")
            return []

    def _generate_content_hash(self, content: str) -> str:
        """Generate SHA-256 hash of content for deduplication."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()[:16]

    async def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics."""
        return {
            "supported_formats": settings.SUPPORTED_FORMATS,
            "max_file_size": settings.MAX_FILE_SIZE,
            "chunk_size": settings.CHUNK_SIZE,
            "chunk_overlap": settings.CHUNK_OVERLAP,
            "max_chunks_per_document": settings.MAX_CHUNKS_PER_DOCUMENT,
            "nlp_available": False,
            "semantic_chunking_available": True
        }
